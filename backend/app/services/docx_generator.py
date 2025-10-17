from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List
import io

class DocxGenerator:
    """DOCX 문서 생성 서비스"""
    
    @staticmethod
    def create_business_plan(
        template_structure: Dict[str, Any],
        generated_content: Dict[str, Any],
        business_info: Dict[str, Any]
    ) -> io.BytesIO:
        """
        사업계획서 DOCX 생성
        
        Args:
            template_structure: 원본 템플릿 구조
            generated_content: AI가 생성한 내용
            business_info: 사업 기본 정보
            
        Returns:
            BytesIO: DOCX 파일 바이트 스트림
        """
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading(business_info.get('title', '사업계획서'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 사업 개요
        if business_info.get('description'):
            doc.add_heading('1. 사업 개요', 1)
            doc.add_paragraph(business_info['description'])
        
        # 시장 분석
        if 'market_analysis' in generated_content:
            doc.add_heading('2. 시장 분석', 1)
            doc.add_paragraph(generated_content['market_analysis'])
        
        # 경쟁사 분석
        if 'competitive_analysis' in generated_content:
            doc.add_heading('3. 경쟁사 분석 및 차별화 전략', 1)
            doc.add_paragraph(generated_content['competitive_analysis'])
        
        # 재무 계획
        if 'financial_plan' in generated_content:
            doc.add_heading('4. 재무 계획', 1)
            doc.add_paragraph(generated_content['financial_plan'])
            
            # 재무 계획 표 추가
            if 'financial_tables' in generated_content:
                for table_data in generated_content['financial_tables']:
                    DocxGenerator._add_table(doc, table_data)
        
        # 추가 섹션들
        if template_structure and 'paragraphs' in template_structure:
            for para in template_structure['paragraphs']:
                section_type = para.get('section_type', 'general')
                if section_type not in ['market_analysis', 'competitive_analysis', 'financial_plan']:
                    # 기타 섹션 추가
                    if para.get('style', '').startswith('Heading'):
                        doc.add_heading(para['text'], 2)
                    else:
                        doc.add_paragraph(para['text'])
        
        # 요구사항 및 주의사항
        if business_info.get('requirements') or business_info.get('notes'):
            doc.add_page_break()
            doc.add_heading('참고사항', 1)
            
            if business_info.get('requirements'):
                doc.add_paragraph('필수 요구사항:', style='Heading 2')
                doc.add_paragraph(business_info['requirements'])
            
            if business_info.get('notes'):
                doc.add_paragraph('주의사항:', style='Heading 2')
                doc.add_paragraph(business_info['notes'])
        
        # BytesIO로 저장
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
    
    @staticmethod
    def _add_table(doc: Document, table_data: Dict[str, Any]):
        """문서에 표 추가"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers or not rows:
            return
        
        # 표 생성
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 헤더 추가
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_cells[idx].text = str(header)
            # 헤더 스타일링
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 데이터 행 추가
        for row_idx, row_data in enumerate(rows, start=1):
            cells = table.rows[row_idx].cells
            for col_idx, cell_data in enumerate(row_data.get('cells', [])):
                if col_idx < len(cells):
                    cells[col_idx].text = str(cell_data.get('value', ''))
        
        doc.add_paragraph()  # 표 아래 여백

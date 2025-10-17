"""
@CODE:docx-parser-service
DOCX 파일 파싱 서비스 (표 분석 특화)

Related:
- @SPEC:FEAT-001 - DOCX Parser with Table Analysis
- @TEST:docx-parser-unit
- @DOC:api-docx-parser
"""

from docx import Document
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

class DocxParser:
    """
    @CODE:docx-parser-service
    DOCX 파일 파싱 서비스 (표 분석 포함)
    
    Implements:
    - @SPEC:FEAT-001-REQ-001 (Document Parsing)
    - @SPEC:FEAT-001-REQ-002 (Table Analysis)
    - @SPEC:FEAT-001-REQ-003 (Section Type Identification)
    - @SPEC:FEAT-001-REQ-004 (Metadata Extraction)
    """
    
    @staticmethod
    def parse_document(file_path: str) -> Dict[str, Any]:
        """
        @CODE:docx-parser-service-parse
        DOCX 문서를 파싱하여 구조화된 데이터 반환
        
        Implements: @SPEC:FEAT-001-REQ-001, @SPEC:FEAT-001-REQ-004
        Tests: @TEST:docx-parser-unit-001, @TEST:docx-parser-unit-004
        
        Args:
            file_path: DOCX 파일 경로
            
        Returns:
            문서 구조 (문단, 표, 메타데이터)
        """
        try:
            doc = Document(file_path)
            
            # 문단 추출
            paragraphs = []
            for idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        "index": idx,
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else "Normal"
                    })
            
            # 표 추출 및 분석 (핵심 기능)
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = DocxParser._parse_table(table, table_idx)
                tables.append(table_data)
            
            # 메타데이터
            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "core_properties": {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or ""
                }
            }
            
            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"DOCX 파싱 오류: {str(e)}")
            raise
    
    @staticmethod
    def _parse_table(table, table_idx: int) -> Dict[str, Any]:
        """
        @CODE:docx-parser-service-table (핵심 기능)
        표를 파싱하여 구조화된 데이터 반환
        
        Implements: @SPEC:FEAT-001-REQ-002
        Tests: @TEST:docx-parser-unit-002
        
        Args:
            table: python-docx Table 객체
            table_idx: 표 인덱스
            
        Returns:
            표 구조 (헤더, 행, 열 정보)
        """
        rows = []
        headers = []
        
        # 첫 번째 행을 헤더로 간주
        if len(table.rows) > 0:
            header_cells = table.rows[0].cells
            headers = [cell.text.strip() for cell in header_cells]
        
        # 데이터 행 추출
        for row_idx, row in enumerate(table.rows[1:], start=1):
            row_data = []
            for cell_idx, cell in enumerate(row.cells):
                row_data.append({
                    "column": headers[cell_idx] if cell_idx < len(headers) else f"Column_{cell_idx}",
                    "value": cell.text.strip()
                })
            rows.append({
                "row_index": row_idx,
                "cells": row_data
            })
        
        return {
            "table_index": table_idx,
            "headers": headers,
            "rows": rows,
            "row_count": len(table.rows) - 1,  # 헤더 제외
            "column_count": len(headers)
        }
    
    @staticmethod
    def identify_section_type(text: str) -> str:
        """
        @CODE:docx-parser-service-section
        문단의 섹션 타입 식별
        
        Implements: @SPEC:FEAT-001-REQ-003
        Tests: @TEST:docx-parser-unit-003
        
        Args:
            text: 문단 텍스트
            
        Returns:
            섹션 타입 (예: "title", "market_analysis", "financial_plan")
        """
        text_lower = text.lower()
        
        # 키워드 기반 섹션 식별
        if any(keyword in text_lower for keyword in ["시장", "분석", "market"]):
            return "market_analysis"
        elif any(keyword in text_lower for keyword in ["경쟁", "competitor", "차별"]):
            return "competitive_analysis"
        elif any(keyword in text_lower for keyword in ["재무", "financial", "손익", "현금"]):
            return "financial_plan"
        elif any(keyword in text_lower for keyword in ["사업", "개요", "summary"]):
            return "business_overview"
        elif any(keyword in text_lower for keyword in ["swot"]):
            return "swot_analysis"
        else:
            return "general"

"""
@TEST:docx-parser-unit
Unit tests for DOCX Parser Service

Related:
- @SPEC:FEAT-001 - DOCX Parser with Table Analysis
- @CODE:docx-parser-service
"""

import pytest
import tempfile
import os
from docx import Document
from backend.app.services.docx_parser import DocxParser


class TestDocxParser:
    """@TEST:docx-parser-unit - DOCX Parser 단위 테스트"""
    
    @pytest.fixture
    def sample_docx(self):
        """테스트용 DOCX 파일 생성"""
        doc = Document()
        doc.add_heading('테스트 사업계획서', 0)
        doc.add_paragraph('이것은 사업 개요입니다.')
        doc.add_paragraph('시장 분석을 진행합니다.')
        
        # 표 추가
        table = doc.add_table(rows=3, cols=3)
        table.rows[0].cells[0].text = '항목'
        table.rows[0].cells[1].text = '2024'
        table.rows[0].cells[2].text = '2025'
        table.rows[1].cells[0].text = '매출'
        table.rows[1].cells[1].text = '100억'
        table.rows[1].cells[2].text = '150억'
        
        # 임시 파일로 저장
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        yield temp_file.name
        
        # 정리
        os.unlink(temp_file.name)
    
    def test_parse_document_success(self, sample_docx):
        """
        @TEST:docx-parser-unit-001
        문서 파싱 성공 테스트
        
        Tests: @SPEC:FEAT-001-REQ-001
        Implements: @CODE:docx-parser-service-parse
        """
        parser = DocxParser()
        result = parser.parse_document(sample_docx)
        
        # Assertions
        assert "paragraphs" in result
        assert "tables" in result
        assert "metadata" in result
        assert len(result["paragraphs"]) >= 2
        assert result["metadata"]["paragraph_count"] >= 2
    
    def test_parse_table_structure(self, sample_docx):
        """
        @TEST:docx-parser-unit-002 (핵심 테스트)
        표 구조 파싱 테스트
        
        Tests: @SPEC:FEAT-001-REQ-002
        Implements: @CODE:docx-parser-service-table
        """
        parser = DocxParser()
        result = parser.parse_document(sample_docx)
        
        # 표가 존재하는지 확인
        assert len(result["tables"]) > 0
        
        table = result["tables"][0]
        
        # 표 구조 검증
        assert "headers" in table
        assert "rows" in table
        assert "table_index" in table
        assert "row_count" in table
        assert "column_count" in table
        
        # 헤더 검증
        assert len(table["headers"]) == 3
        assert "항목" in table["headers"]
        
        # 행 데이터 검증
        assert len(table["rows"]) > 0
        assert table["row_count"] == 2  # 헤더 제외
        assert table["column_count"] == 3
    
    def test_identify_section_type_market_analysis(self):
        """
        @TEST:docx-parser-unit-003
        섹션 타입 식별 테스트 - 시장 분석
        
        Tests: @SPEC:FEAT-001-REQ-003
        Implements: @CODE:docx-parser-service-section
        """
        parser = DocxParser()
        
        # 시장 분석 키워드
        assert parser.identify_section_type("시장 분석") == "market_analysis"
        assert parser.identify_section_type("Market Analysis") == "market_analysis"
        
        # 경쟁사 분석 키워드
        assert parser.identify_section_type("경쟁사 분석") == "competitive_analysis"
        assert parser.identify_section_type("Competitor Analysis") == "competitive_analysis"
        
        # 재무 계획 키워드
        assert parser.identify_section_type("재무 계획") == "financial_plan"
        assert parser.identify_section_type("Financial Plan") == "financial_plan"
        
        # 일반 섹션
        assert parser.identify_section_type("기타 내용") == "general"
    
    def test_metadata_extraction(self, sample_docx):
        """
        @TEST:docx-parser-unit-004
        메타데이터 추출 테스트
        
        Tests: @SPEC:FEAT-001-REQ-004
        Implements: @CODE:docx-parser-service-parse
        """
        parser = DocxParser()
        result = parser.parse_document(sample_docx)
        
        metadata = result["metadata"]
        
        # 메타데이터 구조 검증
        assert "paragraph_count" in metadata
        assert "table_count" in metadata
        assert "core_properties" in metadata
        
        # Core properties 검증
        core_props = metadata["core_properties"]
        assert "title" in core_props
        assert "author" in core_props
        assert "subject" in core_props
        
        # 카운트 검증
        assert metadata["paragraph_count"] > 0
        assert metadata["table_count"] > 0
    
    def test_parse_document_file_not_found(self):
        """
        @TEST:docx-parser-unit-005
        파일 없음 예외 처리 테스트
        
        Tests: Error handling
        """
        parser = DocxParser()
        
        with pytest.raises(Exception):
            parser.parse_document("/nonexistent/file.docx")
    
    def test_empty_table_handling(self):
        """
        @TEST:docx-parser-unit-006
        빈 표 처리 테스트
        
        Tests: Edge case handling
        """
        doc = Document()
        table = doc.add_table(rows=0, cols=0)
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        try:
            parser = DocxParser()
            result = parser.parse_document(temp_file.name)
            
            # 빈 표도 정상 처리되어야 함
            assert "tables" in result
            
        finally:
            os.unlink(temp_file.name)


# Quality Gates (TRUST-5)
# ✅ Test-first: TDD approach
# ✅ Readable: Clear test names and structure
# ✅ Unified: Consistent with pytest patterns
# ✅ Secured: Safe file handling with cleanup
# ✅ Trackable: Full @TAG system

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
